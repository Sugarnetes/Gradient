from dotenv import load_dotenv
from langchain_community.llms.huggingface_hub import HuggingFaceHub
import os
from pdfreader import PdfReader
from docx import Document
from docx.shared import Inches
from docx2pdf import convert

os.environ["HUGGINGFACEHUB_API_TOKEN"] = ""

summarizer = HuggingFaceHub(
    repo_id="mistralai/Mixtral-8x7B-Instruct-v0.1",
    model_kwargs={"temperature": 1, "max_length": 3800}
)


class Summarize():

    def __init__(self, pdf):
        self.pdfTest = PdfReader(pdf)
        self.text_to_return, self.list_of_topics = self.pdfTest.extractor()

    def summarize(self, llm, topic, points):
        return llm(f"For the {topic}, summarize the following bullet points {points} in point format. Do not leave any text out. Make sure the summary is concise. Make sure that the summary is shorter than the original points (do not leave any content out but make sure you reduce any redundancy wherever possible). You can keep on giving larger points than the original - you should be summarizing and giving smaller points.")

    def execute_summary(self):
        doc = Document()
        
        for text in self.text_to_return:
            topic = text[0]
            points = text[1]
            summarized_text = self.summarize(summarizer, topic, points)
            
            # Add the summarized content to the Word document
            doc.add_heading(f"Summary for {topic}", level=1)
            doc.add_paragraph(summarized_text)
            doc.add_paragraph('\n')
        
        # Save the Word document
        doc.save("summarized_content.docx")
        
        # Convert the Word document to PDF
        convert("summarized_content.docx", "summarized_content.pdf")
        os.remove("summarized_content.docx")


if __name__ == "__main__":
    pdfTest = Summarize(r"C:\umer files\Programming PREJE'S\Gradient\server\restapis\denis.pdf")
    pdfTest.execute_summary()
